#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Step 1: 文档转图片
Document to Images Converter

将 PDF、Word、Excel 等文档转换为图片格式。
输入: data/01_raw_documents/
输出: data/02_images/
"""

import os
import sys
import io
import argparse
import logging
from pathlib import Path
from typing import List, Tuple, Optional
from PIL import Image

from config import PATHS, DEFAULT_CONFIG, ensure_directories

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# 检查依赖
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    logger.warning("PyMuPDF 未安装，PDF处理不可用。请运行: pip install pymupdf")

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class DocumentToImageConverter:
    """文档转图片转换器"""

    SUPPORTED_FORMATS = {'.pdf', '.docx', '.doc', '.xlsx', '.xls'}

    def __init__(self, input_dir: Path, output_dir: Path):
        self.input_dir = Path(input_dir)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

        self.dpi = DEFAULT_CONFIG.get("image_dpi", 300)
        self.only_first_page = DEFAULT_CONFIG.get("only_first_page", True)

        self.stats = {
            "total": 0,
            "success": 0,
            "failed": 0,
            "images_created": 0
        }

    def scan_documents(self) -> List[Path]:
        """扫描目录下的所有支持的文档"""
        documents = []
        for ext in self.SUPPORTED_FORMATS:
            documents.extend(self.input_dir.glob(f"**/*{ext}"))
        documents.sort()
        self.stats["total"] = len(documents)
        return documents

    def convert_pdf(self, pdf_path: Path) -> List[Tuple[str, Image.Image]]:
        """将 PDF 转换为图片列表"""
        if not HAS_PYMUPDF:
            raise ImportError("PyMuPDF 未安装")

        images = []
        doc = fitz.open(pdf_path)

        try:
            page_count = doc.page_count
            pages_to_process = 1 if self.only_first_page else page_count

            for page_num in range(pages_to_process):
                page = doc.load_page(page_num)
                # 根据 DPI 计算缩放比例 (72 是 PDF 默认 DPI)
                zoom = self.dpi / 72
                mat = fitz.Matrix(zoom, zoom)
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                image = Image.open(io.BytesIO(img_data))
                images.append((f"page_{page_num + 1}", image))

        finally:
            doc.close()

        return images

    def convert_docx(self, docx_path: Path) -> List[Tuple[str, Image.Image]]:
        """将 Word 文档转换为图片（通过 LibreOffice 或文本渲染）"""
        # 尝试使用 LibreOffice 转换
        try:
            import subprocess
            import tempfile

            with tempfile.TemporaryDirectory() as tmp_dir:
                # 先转为 PDF
                cmd = [
                    'libreoffice', '--headless', '--convert-to', 'pdf',
                    '--outdir', tmp_dir, str(docx_path)
                ]
                result = subprocess.run(cmd, capture_output=True, timeout=60)

                pdf_path = Path(tmp_dir) / f"{docx_path.stem}.pdf"
                if pdf_path.exists():
                    return self.convert_pdf(pdf_path)

        except Exception as e:
            logger.warning(f"LibreOffice 转换失败: {e}")

        # 降级方案：提取文本并渲染为图片
        if HAS_DOCX:
            return self._render_docx_as_image(docx_path)

        raise Exception("无法转换 Word 文档")

    def _render_docx_as_image(self, docx_path: Path) -> List[Tuple[str, Image.Image]]:
        """将 Word 文档文本渲染为图片"""
        from PIL import ImageDraw, ImageFont

        doc = Document(docx_path)
        text_lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_lines.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_lines.append(row_text)

        # 创建图片
        font_size = 14
        line_height = font_size + 6
        padding = 20
        max_width = 1200

        try:
            font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", font_size)
        except:
            font = ImageFont.load_default()

        height = len(text_lines) * line_height + padding * 2
        image = Image.new('RGB', (max_width, max(height, 100)), 'white')
        draw = ImageDraw.Draw(image)

        y = padding
        for line in text_lines:
            draw.text((padding, y), line[:150], fill='black', font=font)
            y += line_height

        return [("page_1", image)]

    def convert_excel(self, excel_path: Path) -> List[Tuple[str, Image.Image]]:
        """将 Excel 转换为图片"""
        # 同样尝试 LibreOffice
        try:
            import subprocess
            import tempfile

            with tempfile.TemporaryDirectory() as tmp_dir:
                cmd = [
                    'libreoffice', '--headless', '--convert-to', 'pdf',
                    '--outdir', tmp_dir, str(excel_path)
                ]
                subprocess.run(cmd, capture_output=True, timeout=60)

                pdf_path = Path(tmp_dir) / f"{excel_path.stem}.pdf"
                if pdf_path.exists():
                    return self.convert_pdf(pdf_path)

        except Exception as e:
            logger.warning(f"LibreOffice 转换失败: {e}")

        # 降级方案：渲染为文本图片
        return self._render_excel_as_image(excel_path)

    def _render_excel_as_image(self, excel_path: Path) -> List[Tuple[str, Image.Image]]:
        """将 Excel 内容渲染为图片"""
        from PIL import ImageDraw, ImageFont

        if not HAS_OPENPYXL:
            raise ImportError("openpyxl 未安装")

        wb = openpyxl.load_workbook(excel_path)
        text_lines = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            text_lines.append(f"=== {sheet_name} ===")
            for row in ws.iter_rows(values_only=True, max_row=50):
                row_text = " | ".join(str(c) if c else "" for c in row)
                if row_text.strip():
                    text_lines.append(row_text[:200])

        # 渲染
        font_size = 12
        line_height = font_size + 4
        padding = 20
        max_width = 1400

        try:
            font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", font_size)
        except:
            font = ImageFont.load_default()

        height = len(text_lines) * line_height + padding * 2
        image = Image.new('RGB', (max_width, max(height, 100)), 'white')
        draw = ImageDraw.Draw(image)

        y = padding
        for line in text_lines:
            draw.text((padding, y), line, fill='black', font=font)
            y += line_height

        return [("page_1", image)]

    def convert_document(self, doc_path: Path) -> List[Tuple[str, Image.Image]]:
        """根据文件类型转换文档"""
        ext = doc_path.suffix.lower()

        if ext == '.pdf':
            return self.convert_pdf(doc_path)
        elif ext in ['.docx', '.doc']:
            return self.convert_docx(doc_path)
        elif ext in ['.xlsx', '.xls']:
            return self.convert_excel(doc_path)
        else:
            raise ValueError(f"不支持的格式: {ext}")

    def process_single(self, doc_path: Path) -> bool:
        """处理单个文档"""
        try:
            logger.info(f"处理: {doc_path.name}")
            images = self.convert_document(doc_path)

            if not images:
                logger.warning(f"  未生成图片: {doc_path.name}")
                return False

            # 保存图片
            for page_name, image in images:
                # 文件名格式: 原文件名_页码.png
                output_name = f"{doc_path.stem}_{page_name}.png"
                output_path = self.output_dir / output_name
                image.save(output_path, "PNG")
                logger.info(f"  ✅ 保存: {output_name}")
                self.stats["images_created"] += 1

            self.stats["success"] += 1
            return True

        except Exception as e:
            logger.error(f"  ❌ 失败: {doc_path.name} - {e}")
            self.stats["failed"] += 1
            return False

    def run(self):
        """运行转换"""
        logger.info("=" * 60)
        logger.info("Step 1: 文档转图片")
        logger.info("=" * 60)
        logger.info(f"输入目录: {self.input_dir}")
        logger.info(f"输出目录: {self.output_dir}")

        documents = self.scan_documents()
        if not documents:
            logger.warning("未找到任何支持的文档文件")
            return self.stats

        logger.info(f"找到 {len(documents)} 个文档\n")

        for i, doc_path in enumerate(documents, 1):
            logger.info(f"[{i}/{len(documents)}]")
            self.process_single(doc_path)

        # 打印统计
        logger.info("\n" + "=" * 60)
        logger.info("转换完成")
        logger.info("=" * 60)
        logger.info(f"总文档数: {self.stats['total']}")
        logger.info(f"成功: {self.stats['success']}")
        logger.info(f"失败: {self.stats['failed']}")
        logger.info(f"生成图片数: {self.stats['images_created']}")

        return self.stats


def main():
    parser = argparse.ArgumentParser(description='Step 1: 文档转图片')
    parser.add_argument('-i', '--input', type=str, help='输入目录')
    parser.add_argument('-o', '--output', type=str, help='输出目录')
    args = parser.parse_args()

    ensure_directories()

    input_dir = Path(args.input) if args.input else PATHS["input_documents"]
    output_dir = Path(args.output) if args.output else PATHS["step1_images"]

    converter = DocumentToImageConverter(input_dir, output_dir)
    converter.run()


if __name__ == "__main__":
    main()
