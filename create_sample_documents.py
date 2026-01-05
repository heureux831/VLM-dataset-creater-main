#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建示例文档脚本
Create Sample Documents Script

此脚本创建用于测试多格式文档支持的示例文档
"""

import os
import sys
from pathlib import Path

def create_sample_excel():
    """创建示例Excel文档"""
    try:
        import openpyxl
    except ImportError:
        print("❌ 需要安装 openpyxl: pip install openpyxl")
        return False

    wb = openpyxl.Workbook()

    # 工作表1: 海运单信息
    ws1 = wb.active
    ws1.title = "海运单信息"
    ws1.append(["字段名", "值"])
    ws1.append(["托运人", "ABC Import & Export Co., Ltd."])
    ws1.append(["收货人", "XYZ Trading Company"])
    ws1.append(["通知方", "DEF Logistics Ltd."])
    ws1.append(["装货港", "Shanghai, China"])
    ws1.append(["卸货港", "Los Angeles, USA"])
    ws1.append(["交货港", "Long Beach, USA"])
    ws1.append(["交货地点", "Long Beach Terminal"])
    ws1.append(["收货地点", "Shanghai Warehouse"])
    ws1.append(["船名", "MV Ocean Star"])
    ws1.append(["航次", "Voyage No. 2024-001"])
    ws1.append(["集装箱号", "MSKU1234567"])
    ws1.append(["封号", "SEAL123456"])
    ws1.append(["提单号", "BOL20240001"])

    # 工作表2: 货物信息
    ws2 = wb.create_sheet("货物信息")
    ws2.append(["包装", "货物描述", "重量", "体积"])
    ws2.append(["100 Cartons", "Electronic Components", "500 KG", "2.5 CBM"])
    ws2.append(["50 Boxes", "Textile Products", "300 KG", "1.8 CBM"])
    ws2.append(["20 Pallets", "Machinery Parts", "1000 KG", "5.0 CBM"])

    # 工作表3: 费用信息
    ws3 = wb.create_sheet("费用信息")
    ws3.append(["项目", "费率", "总计"])
    ws3.append(["海运费", "$50/CBM", "$450"])
    ws3.append(["附加费", "$100/TEU", "$100"])
    ws3.append(["总计", "", "$550"])

    filename = "sample_bill_of_lading.xlsx"
    wb.save(filename)
    print(f"✅ 已创建示例Excel文档: {filename}")
    return True


def create_sample_word():
    """创建示例Word文档"""
    try:
        from docx import Document
    except ImportError:
        print("❌ 需要安装 python-docx: pip install python-docx")
        return False

    doc = Document()

    # 添加标题
    doc.add_heading('海运单 (BILL OF LADING)', 0)

    # 添加内容
    doc.add_heading('托运人 (Shipper):', level=1)
    doc.add_paragraph('ABC Import & Export Co., Ltd.')
    doc.add_paragraph('Address: 123 Shanghai Road, China')
    doc.add_paragraph('Phone: +86-21-12345678')

    doc.add_heading('收货人 (Consignee):', level=1)
    doc.add_paragraph('XYZ Trading Company')
    doc.add_paragraph('Address: 456 Los Angeles St., USA')
    doc.add_paragraph('Phone: +1-310-87654321')

    doc.add_heading('通知方 (Notify Party):', level=1)
    doc.add_paragraph('DEF Logistics Ltd.')
    doc.add_paragraph('Address: 789 Long Beach Ave., USA')

    doc.add_heading('运输信息 (Transport Information):', level=1)
    doc.add_paragraph('船名 (Vessel): MV Ocean Star')
    doc.add_paragraph('航次 (Voyage): Voyage No. 2024-001')
    doc.add_paragraph('装货港 (Port of Loading): Shanghai, China')
    doc.add_paragraph('卸货港 (Port of Discharge): Los Angeles, USA')
    doc.add_paragraph('交货港 (Port of Delivery): Long Beach, USA')

    doc.add_heading('集装箱信息 (Container Information):', level=1)
    doc.add_paragraph('集装箱号 (Container No.): MSKU1234567')
    doc.add_paragraph('封号 (Seal No.): SEAL123456')

    doc.add_heading('提单信息 (B/L Information):', level=1)
    doc.add_paragraph('提单号 (B/L No.): BOL20240001')
    doc.add_paragraph('日期 (Date): 2024-01-05')

    doc.add_heading('货物信息 (Cargo Information):', level=1)
    doc.add_paragraph('货物描述 (Description of Goods): Electronic Components')
    doc.add_paragraph('包装 (Package): 100 Cartons')
    doc.add_paragraph('重量 (Weight): 500 KG')
    doc.add_paragraph('体积 (Volume): 2.5 CBM')

    doc.add_heading('费用信息 (Freight Information):', level=1)
    doc.add_paragraph('海运费 (Freight): $450')
    doc.add_paragraph('总计 (Total): $550')

    filename = "sample_bill_of_lading.docx"
    doc.save(filename)
    print(f"✅ 已创建示例Word文档: {filename}")
    return True


def create_sample_pdf():
    """创建示例PDF（需要用户手动创建）"""
    pdf_note = """
# 示例PDF文档创建说明

由于需要专门的库来创建PDF，建议您：

1. 使用现有的海运单PDF文档
2. 或者从网上下载海运单样本
3. 或者将Word/Excel文档另存为PDF

推荐的PDF文档应包含以下信息：
- 托运人、收货人、通知方
- 港口信息（装货港、卸货港等）
- 船名、航次
- 集装箱号、封号
- 提单号
- 货物描述、包装、重量、体积
- 费用信息

请将PDF文件命名为：sample_bill_of_lading.pdf
"""
    print(pdf_note)

    with open("CREATE_PDF_NOTE.md", "w", encoding="utf-8") as f:
        f.write(pdf_note)

    print("✅ 已创建PDF创建说明: CREATE_PDF_NOTE.md")
    return True


def main():
    """主函数"""
    print("\n" + "="*60)
    print("创建示例文档")
    print("="*60)
    print()

    # 创建输入目录
    input_dir = Path("bills_of_lading")
    input_dir.mkdir(exist_ok=True)
    print(f"✅ 输入目录已准备: {input_dir}")
    print()

    # 创建示例文档
    success_count = 0

    print("1. 创建示例Excel文档...")
    if create_sample_excel():
        success_count += 1

    print("\n2. 创建示例Word文档...")
    if create_sample_word():
        success_count += 1

    print("\n3. 创建示例PDF说明...")
    if create_sample_pdf():
        success_count += 1

    # 总结
    print("\n" + "="*60)
    print("创建完成")
    print("="*60)
    print(f"✅ 成功创建 {success_count}/3 个示例文档")

    print("\n📁 文件列表:")
    for file in Path(".").glob("sample_*"):
        print(f"   - {file.name}")

    print("\n📝 下一步操作:")
    print("1. 将创建的示例文档移动到输入目录:")
    print("   mv sample_bill_of_lading.xlsx bills_of_lading/")
    print("   mv sample_bill_of_lading.docx bills_of_lading/")
    print()
    print("2. 添加您的PDF文档到输入目录:")
    print("   cp your_bill_of_lading.pdf bills_of_lading/")
    print()
    print("3. 运行多格式文档处理:")
    print("   python multi_format_to_funsd.py")
    print()
    print("或运行配置检查:")
    print("   python check_config.py")
    print()

    return 0 if success_count > 0 else 1


if __name__ == "__main__":
    try:
        sys.exit(main())
    except KeyboardInterrupt:
        print("\n\n⚠️  用户中断操作")
        sys.exit(0)
    except Exception as e:
        print(f"\n❌ 发生错误: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
