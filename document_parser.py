#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
多格式文档解析器 - 支持PDF、Excel、Word文档
Multi-Format Document Parser - Support PDF, Excel, Word documents

此模块提供统一的接口来处理不同格式的文档，
将其转换为图像后使用PaddleOCR进行文本识别。
"""

import os
import sys
import io
from pathlib import Path
from typing import List, Tuple, Dict, Any, Optional
import logging

# 尝试导入各种文档处理库
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from docx2pdf import convert
    HAS_DOCX2PDF = True
except ImportError:
    HAS_DOCX2PDF = False

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class DocumentParser:
    """多格式文档解析器"""

    SUPPORTED_FORMATS = {
        '.pdf': 'PDF文档',
        '.xlsx': 'Excel工作簿',
        '.xls': 'Excel工作簿',
        '.docx': 'Word文档',
        '.doc': 'Word文档'
    }

    def __init__(self, temp_dir: str = "./temp_docs"):
        """
        初始化文档解析器

        Args:
            temp_dir: 临时文件存储目录
        """
        self.temp_dir = Path(temp_dir)
        self.temp_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"临时目录: {self.temp_dir}")

    def parse_document(self, file_path: str) -> List[Tuple[str, Image.Image]]:
        """
        解析文档并返回图像列表

        Args:
            file_path: 文档文件路径

        Returns:
            List of (page_text, image) tuples

        Raises:
            ValueError: 不支持的文档格式
            FileNotFoundError: 文件不存在
            Exception: 解析过程中出现错误
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        file_ext = file_path.suffix.lower()

        if file_ext not in self.SUPPORTED_FORMATS:
            raise ValueError(f"不支持的文档格式: {file_ext}")

        logger.info(f"开始解析 {self.SUPPORTED_FORMATS[file_ext]}: {file_path.name}")

        if file_ext == '.pdf':
            return self._parse_pdf(file_path)
        elif file_ext in ['.xlsx', '.xls']:
            return self._parse_excel(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self._parse_word(file_path)
        else:
            raise ValueError(f"未实现的文档格式处理: {file_ext}")

    def _parse_pdf(self, pdf_path: Path) -> List[Tuple[str, Image.Image]]:
        """解析PDF文档"""
        if not HAS_PYMUPDF:
            raise ImportError("PyMuPDF (fitz) 未安装，请运行: pip install pymupdf")

        images = []
        try:
            doc = fitz.open(pdf_path)
            logger.info(f"PDF总页数: {len(doc)}")

            for page_num in range(len(doc)):
                page = doc.load_page(page_num)

                # 转换为图像
                mat = fitz.Matrix(2.0, 2.0)  # 2倍缩放提高OCR准确率
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")

                # 使用PIL转换为Image对象
                if HAS_PIL:
                    image = Image.open(io.BytesIO(img_data))
                    page_text = page.get_text()
                    images.append((page_text, image))
                    logger.info(f"已处理PDF第 {page_num + 1} 页")
                else:
                    raise ImportError("PIL未安装，请运行: pip install Pillow")

            doc.close()
            logger.info(f"PDF解析完成，共 {len(images)} 页")
            return images

        except Exception as e:
            logger.error(f"PDF解析错误: {str(e)}")
            raise

    def _parse_excel(self, excel_path: Path) -> List[Tuple[str, Image.Image]]:
        """解析Excel文档"""
        if not HAS_OPENPYXL:
            raise ImportError("openpyxl未安装，请运行: pip install openpyxl")

        if not HAS_PIL:
            raise ImportError("PIL未安装，请运行: pip install Pillow")

        images = []
        try:
            # 加载Excel文件
            wb = openpyxl.load_workbook(excel_path)
            logger.info(f"Excel工作表: {wb.sheetnames}")

            all_text = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                sheet_text = f"\n=== 工作表: {sheet_name} ===\n"

                for row in ws.iter_rows(values_only=True):
                    row_text = []
                    for cell in row:
                        if cell is not None:
                            row_text.append(str(cell))
                    if row_text:
                        sheet_text += " | ".join(row_text) + "\n"

                all_text.append(sheet_text)

            # 合并所有工作表的文本
            full_text = "\n".join(all_text)

            # 创建图像显示Excel内容
            # 使用简单的文本渲染
            from PIL import Image, ImageDraw, ImageFont

            # 计算图像大小
            lines = full_text.split('\n')
            max_line_length = max(len(line) for line in lines) if lines else 0
            width = min(max_line_length * 10, 2000)  # 限制最大宽度
            height = len(lines) * 20 + 50

            # 创建图像
            image = Image.new('RGB', (width, height), 'white')
            draw = ImageDraw.Draw(image)

            # 尝试使用默认字体
            try:
                font = ImageFont.truetype("arial.ttf", 12)
            except:
                font = ImageFont.load_default()

            # 绘制文本
            y_position = 10
            for line in lines:
                if line.strip():
                    draw.text((10, y_position), line, fill='black', font=font)
                    y_position += 18

            images.append((full_text, image))
            logger.info(f"Excel解析完成，共 {len(images)} 个图像")

            return images

        except Exception as e:
            logger.error(f"Excel解析错误: {str(e)}")
            raise

    def _parse_word(self, word_path: Path) -> List[Tuple[str, Image.Image]]:
        """解析Word文档"""
        # 方法1: 直接解析docx内容
        if word_path.suffix.lower() == '.docx' and HAS_DOCX:
            try:
                doc = Document(word_path)
                full_text = []

                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text)

                # 解析表格
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            full_text.append(" | ".join(row_text))

                text_content = "\n".join(full_text)

                # 创建图像显示Word内容
                if HAS_PIL:
                    from PIL import Image, ImageDraw, ImageFont

                    lines = text_content.split('\n')
                    max_line_length = max(len(line) for line in lines) if lines else 0
                    width = min(max_line_length * 10, 2000)
                    height = len(lines) * 20 + 50

                    image = Image.new('RGB', (width, height), 'white')
                    draw = ImageDraw.Draw(image)

                    try:
                        font = ImageFont.truetype("arial.ttf", 12)
                    except:
                        font = ImageFont.load_default()

                    y_position = 10
                    for line in lines:
                        if line.strip():
                            draw.text((10, y_position), line, fill='black', font=font)
                            y_position += 18

                    logger.info(f"Word解析完成(docx)，共1个图像")
                    return [(text_content, image)]
                else:
                    raise ImportError("PIL未安装")

            except Exception as e:
                logger.error(f"Word(docx)解析错误: {str(e)}")
                raise

        # 方法2: 转换为PDF再解析
        elif word_path.suffix.lower() == '.doc' or (word_path.suffix.lower() == '.docx' and not HAS_DOCX):
            if HAS_DOCX2PDF:
                try:
                    logger.info("将Word文档转换为PDF...")
                    pdf_path = self.temp_dir / f"{word_path.stem}.pdf"
                    convert(str(word_path), str(pdf_path))

                    # 解析生成的PDF
                    result = self._parse_pdf(pdf_path)

                    # 清理临时PDF文件
                    pdf_path.unlink()

                    return result

                except Exception as e:
                    logger.error(f"Word转PDF失败: {str(e)}")
                    raise Exception(f"无法处理.doc文件，请转换为.docx格式或安装python-docx库")
            else:
                raise ImportError(
                    "处理.doc文件需要安装docx2pdf，请运行: pip install docx2pdf\n"
                    "或将文件转换为.docx格式"
                )

        else:
            raise ValueError(f"不支持的Word文档格式: {word_path.suffix}")

    def cleanup(self):
        """清理临时文件"""
        import shutil
        if self.temp_dir.exists():
            shutil.rmtree(self.temp_dir)
            logger.info("临时文件已清理")


def main():
    """测试函数"""
    parser = DocumentParser()

    test_file = input("请输入文档文件路径: ").strip()

    try:
        result = parser.parse_document(test_file)
        print(f"\n解析成功！共 {len(result)} 页/个图像")

        for i, (text, image) in enumerate(result):
            print(f"\n--- 第 {i + 1} 页 ---")
            print(f"文本长度: {len(text)} 字符")
            print(f"图像大小: {image.size}")

            # 保存图像
            output_path = f"output_page_{i + 1}.png"
            image.save(output_path)
            print(f"图像已保存: {output_path}")

    except Exception as e:
        print(f"解析失败: {str(e)}")
        import traceback
        traceback.print_exc()

    finally:
        parser.cleanup()


if __name__ == "__main__":
    main()
